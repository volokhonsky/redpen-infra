#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script to create .docx documents from book page images, organized by paragraphs.

This script processes a collection of book page images and creates .docx documents
for each paragraph or section as defined in the table of contents. The script:
1. Parses the table of contents to extract paragraph information (title, start page, end page)
2. For each paragraph, creates a .docx document containing the images of the pages in that paragraph
3. Saves the documents with appropriate filenames based on their content

Usage:
    ./create_paragraph_docs.py

Requirements:
    - python-docx: For creating Word documents
    - Pillow: For image processing

Note:
    - The page numbering in the images is offset by 1 (page_004.png is actually page 3)
    - The script expects the images to be named in the format page_XXX.png
"""

import os
import re
from docx import Document
from docx.shared import Inches
from PIL import Image
import shutil
import tempfile

# Define the path to the images directory
IMAGES_DIR = "redpen-content/medinsky11klass/images"
# Define the path to save the output documents
OUTPUT_DIR = "redpen-content/medinsky11klass/paragraphs"
# Define the offset between image numbers and actual page numbers
PAGE_OFFSET = 1

# Create the output directory if it doesn't exist
os.makedirs(OUTPUT_DIR, exist_ok=True)

def parse_toc():
    """
    Parse the table of contents to extract paragraph information.
    Returns a list of dictionaries, each containing:
    - title: The title of the paragraph
    - start_page: The starting page number
    - end_page: The ending page number (or None if it's the last paragraph in a section)
    """
    # The table of contents as provided in the issue description
    toc_text = """
Введение . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . 3
Глава I. СССР в 1945—1991 гг. . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . 5
§ 1. Восстановление и развитие экономики и социальной сферы . . . . . 6
§ 2. Политическая система в послевоенные годы . . . . . . . . . . . . . . . . . . . 21
§ 3. Идеология, наука, культура и спорт в послевоенные годы . . . . . . . 29
§ 4. Место и роль СССР в послевоенном мире. 
Внешняя политика СССР в 1945—1953 гг. . . . . . . . . . . . . . . . . . . . . . . 38
§ 5. Новое руководство страны. Смена политического курса . . . . . . . . . 59
§ 6. Экономическое и социальное развитие в 1953—1964 гг. . . . . . . . . . . 72
§ 7. Развитие науки и техники в СССР в 1953—1964 гг. . . . . . . . . . . . . . 82
§ 8. Культурное пространство в 1953—1964 гг. . . . . . . . . . . . . . . . . . . . . . 94
§ 9. Перемены в повседневной жизни в 1953—1964 гг. . . . . . . . . . . . . . . 106
§ 10. Внешняя политика в 1953—1964 гг. . . . . . . . . . . . . . . . . . . . . . . . . . . . . 122
§ 11. Политическое развитие СССР в 1964—1985 гг. . . . . . . . . . . . . . . . . . 134
§ 12. Социально-экономическое развитие СССР в 1964—1985 гг. . . . . . 144
§ 13. Развитие науки, образования, здравоохранения 
в 1964—1985 гг. . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . 154
§ 14. Идеология и культура в 1964—1985 гг. . . . . . . . . . . . . . . . . . . . . . . . . . 164
§ 15. Повседневная жизнь советского общества в 1964—1985 гг. . . . . . . . 177
§ 16. Национальная политика и национальные движения 
в 1964—1985 гг. . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . 184
§ 17. Внешняя политика СССР в 1964—1985 гг. . . . . . . . . . . . . . . . . . . . . . 192
§ 18. СССР и мир в начале 1980-х гг. Предпосылки реформ . . . . . . . . . . 202
§ 19. Социально-экономическое развитие СССР 
в 1985—1991 гг. . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . 210
§ 20. Перемены в духовной сфере в годы перестройки . . . . . . . . . . . . . . . 224
§ 21. Реформа политической системы СССР и её итоги . . . . . . . . . . . . . . 234
§ 22. Новое политическое мышление и перемены 
во внешней политике . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . 243
§ 23. Национальная политика и подъём национальных 
движений. Распад СССР . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . 252
Итоги главы . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . 269
Вопросы и задания к главе . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . 270
Темы проектов . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . 270
Ресурсы к главе . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . 271
447 Оглавление
Глава II. Российская Федерация в 1992 — начале 2020-х гг. . . . . . . . . . 277
§ 24. Российская экономика в условиях рынка . . . . . . . . . . . . . . . . . . . . . . 278
§ 25. Политическое развитие Российской Федерации 
в 1990-е гг. . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . 293
§ 26. Межнациональные отношения и национальная 
политика в 1990-е гг. . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . 303
§ 27. Повседневная жизнь в 1990-е гг. . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . 308
§ 28. Россия и мир. Внешняя политика Российской Федерации 
в 1990-е гг. . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . 313
§ 29. Политические вызовы и новые приоритеты внутренней 
политики России в начале XXI в. . . . . . . . . . . . . . . . . . . . . . . . . . . . . . 325
§ 30. Россия в 2008—2011 гг. . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . 335
§ 31. Социально-экономическое развитие России в начале XXI в. 
Приоритетные национальные проекты . . . . . . . . . . . . . . . . . . . . . . . . 340
§ 32—33. Культура, наука, спорт и общественная жизнь
в 1990-х — начале 2020-х гг. . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . 348
§ 34—35. Внешняя политика в начале XXI в. Россия 
в современном мире . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . 372
§ 36. Россия в 2012 — начале 2020-х гг. . . . . . . . . . . . . . . . . . . . . . . . . . . . . . 382
§ 37. Россия сегодня. Специальная военная операция (СВО) . . . . . . . . . 390
Вопросы и задания к главе . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . 419
    """

    # Regular expression to match paragraph entries
    # Format: § <number>. <title> . . . <page>
    # or other entries like "Введение", "Глава", etc.
    pattern = r'(§\s+\d+(?:—\d+)?\.|\w+)\s+(.*?)(?:\s+\.+\s+)(\d+)$'

    # First, preprocess the lines to handle multi-line entries
    lines = toc_text.strip().split('\n')

    # Process multi-line entries
    processed_lines = []
    current_line = ""
    current_prefix = ""

    for line in lines:
        # Skip the "Оглавление" line
        if "Оглавление" in line and not re.search(pattern, line):
            continue

        # Check if this is a new entry
        match = re.search(pattern, line)
        if match:
            if current_line:
                processed_lines.append(current_line)
            current_line = line
            current_prefix = match.group(1)
        elif line.strip() and current_line:
            # This might be a continuation of the previous line
            # Check if it's a continuation or a new entry without a page number yet
            if re.match(r'^[§\w]', line.strip()):
                # This looks like a new entry, but it might be incomplete
                # Store the previous line
                if current_line:
                    processed_lines.append(current_line)
                current_line = line
            else:
                # This is a continuation of the previous line
                current_line += " " + line.strip()
        elif line.strip():
            # This is a new line but not matching our pattern
            current_line = line

    if current_line:
        processed_lines.append(current_line)

    # Now process the lines to extract paragraph information
    paragraphs = []

    for i, line in enumerate(processed_lines):
        match = re.search(pattern, line)
        if match:
            prefix, title, page = match.groups()

            # Clean up the title
            full_title = f"{prefix} {title}".strip()
            start_page = int(page)

            # Try to find the next paragraph to determine the end page
            end_page = None
            for j in range(i + 1, len(processed_lines)):
                next_match = re.search(pattern, processed_lines[j])
                if next_match:
                    end_page = int(next_match.group(3)) - 1
                    break

            paragraphs.append({
                'title': full_title,
                'start_page': start_page,
                'end_page': end_page
            })

    return paragraphs

def create_docx_from_images(paragraph_info, all_paragraphs):
    """
    Create a .docx document from a group of images.

    Args:
        paragraph_info: A dictionary containing paragraph information
            - title: The title of the paragraph
            - start_page: The starting page number
            - end_page: The ending page number (or None if it's the last paragraph in a section)
        all_paragraphs: A list of all paragraphs from the table of contents

    Returns:
        The path to the created document
    """
    title = paragraph_info['title']
    start_page = paragraph_info['start_page']
    end_page = paragraph_info['end_page']

    # If end_page is None, use the next paragraph's start_page - 1 or a large number
    if end_page is None:
        end_page = 1000  # A large number to include all remaining pages

    # Create a new document
    doc = Document()
    doc.add_heading(title, level=1)

    # Create a temporary directory for image processing
    with tempfile.TemporaryDirectory() as temp_dir:
        # Add each page as an image
        for page_num in range(start_page, end_page + 1):
            # Calculate the image file number (with offset)
            image_num = page_num + PAGE_OFFSET
            image_path = os.path.join(IMAGES_DIR, f"page_{image_num:03d}.png")

            if os.path.exists(image_path):
                # Open the image to get its dimensions
                with Image.open(image_path) as img:
                    width, height = img.size

                # Add page marker before the image
                doc.add_paragraph(f"[[page::{page_num}]]")

                # Add the image to the document
                # Scale the image to fit within the page width
                doc.add_picture(image_path, width=Inches(6))
            else:
                print(f"Warning: Image for page {page_num} (file: {image_path}) not found.")

        # Create a safe filename based on paragraph number
        # First, try to extract the paragraph number from the title
        para_match = re.match(r'§\s+(\d+(?:—\d+)?)', title)

        # For paragraphs with clear numbering
        if para_match:
            para_num = para_match.group(1)
            output_path = os.path.join(OUTPUT_DIR, f"para_{para_num}.docx")
        # For special sections
        elif title.startswith("Глава"):
            output_path = os.path.join(OUTPUT_DIR, f"chapter_{len(title) > 6 and title[6] or ''}.docx")
        elif title.startswith("Введение"):
            output_path = os.path.join(OUTPUT_DIR, "intro.docx")
        elif title.startswith("Итоги"):
            output_path = os.path.join(OUTPUT_DIR, "summary.docx")
        elif title.startswith("Вопросы"):
            output_path = os.path.join(OUTPUT_DIR, "questions.docx")
        elif title.startswith("Темы"):
            output_path = os.path.join(OUTPUT_DIR, "topics.docx")
        elif title.startswith("Ресурсы"):
            output_path = os.path.join(OUTPUT_DIR, "resources.docx")
        else:
            # For multi-line paragraph entries or continuations, we need to determine which paragraph they belong to

            # Get the start page of this paragraph
            page_num = paragraph_info['start_page']

            # First, check if this is a continuation of a multi-line paragraph title
            # Look for paragraphs that have this title as part of their full title
            found_para = False
            for i, para in enumerate(all_paragraphs):
                if 'title' in para and para['title'].startswith('§'):
                    para_title = para['title']
                    # Check if this paragraph's title contains our title (for multi-line titles)
                    if title in para_title:
                        para_match = re.match(r'§\s+(\d+(?:—\d+)?)', para_title)
                        if para_match:
                            para_num = para_match.group(1)
                            output_path = os.path.join(OUTPUT_DIR, f"para_{para_num}.docx")
                            found_para = True
                            break

            # If not found as part of a title, check if it's a page within a paragraph's range
            if not found_para:
                for i, para in enumerate(all_paragraphs):
                    if 'title' in para and para['title'].startswith('§'):
                        para_title = para['title']
                        para_match = re.match(r'§\s+(\d+(?:—\d+)?)', para_title)
                        if para_match and para['start_page'] <= page_num and (para['end_page'] is None or para['end_page'] >= page_num):
                            para_num = para_match.group(1)
                            output_path = os.path.join(OUTPUT_DIR, f"para_{para_num}.docx")
                            found_para = True
                            break

            # If we still couldn't find a matching paragraph, use a cleaned version of the title
            if not found_para:
                safe_title = re.sub(r'[^\w\s-]', '', title)
                safe_title = re.sub(r'[\s-]+', '_', safe_title)
                safe_title = safe_title.lstrip('_')
                output_path = os.path.join(OUTPUT_DIR, f"{safe_title}.docx")

        # Save the document
        doc.save(output_path)

    return output_path

def main():
    """Main function to orchestrate the process."""
    # Clean up the output directory
    if os.path.exists(OUTPUT_DIR):
        print(f"Cleaning up output directory: {OUTPUT_DIR}")
        for file in os.listdir(OUTPUT_DIR):
            file_path = os.path.join(OUTPUT_DIR, file)
            if os.path.isfile(file_path):
                os.remove(file_path)

    print("Parsing table of contents...")
    paragraphs = parse_toc()

    print(f"Found {len(paragraphs)} paragraphs.")

    # Process each paragraph
    for i, paragraph in enumerate(paragraphs):
        print(f"Processing {i+1}/{len(paragraphs)}: {paragraph['title']}")
        print(f"  Pages: {paragraph['start_page']} to {paragraph['end_page'] or 'end'}")

        # Create a .docx document for this paragraph
        output_path = create_docx_from_images(paragraph, paragraphs)
        print(f"  Created: {output_path}")

if __name__ == "__main__":
    main()
