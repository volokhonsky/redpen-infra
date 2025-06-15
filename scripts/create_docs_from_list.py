#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script to create .docx documents from book page images, organized by paragraphs.

This script reads a list of paragraphs from a file and creates .docx documents
for each paragraph, including the images of the pages in that paragraph.

Usage:
    ./create_docs_from_list.py

Requirements:
    - python-docx: For creating Word documents
    - Pillow: For image processing

Note:
    - The page numbering in the images is offset by 1 (page_004.png is actually page 3)
    - The script expects the images to be named in the format page_XXX.png
"""

import os
import re
from docx import Document
from docx.shared import Inches
from PIL import Image
import tempfile

# Define the path to the images directory - using images_with_grid instead of images
IMAGES_DIR = "redpen-content/medinsky11klass/images_with_grid"
# Define the path to save the output documents
OUTPUT_DIR = "redpen-content/medinsky11klass/paragraphs"
# Define the path to the paragraphs list file
PARAGRAPHS_LIST_FILE = "redpen-content/medinsky11klass/paragraphs_list.txt"
# Define the offset between image numbers and actual page numbers
PAGE_OFFSET = 1

# Create the output directory if it doesn't exist
os.makedirs(OUTPUT_DIR, exist_ok=True)

def read_paragraphs_list():
    """
    Read the paragraphs list from the file.
    Returns a list of dictionaries, each containing:
    - id: The paragraph ID (number or code)
    - title: The title of the paragraph
    - start_page: The starting page number
    - end_page: The ending page number
    """
    paragraphs = []

    with open(PARAGRAPHS_LIST_FILE, 'r', encoding='utf-8') as f:
        for line in f:
            line = line.strip()
            # Skip empty lines, comments, and section headers
            if not line or line.startswith('#'):
                continue

            # Parse the line: id, title, start_page, end_page
            # Use regex to handle commas within the title
            match = re.match(r'^([^,]+),\s*(.+),\s*(\d+),\s*(\d+)$', line)
            if match:
                para_id = match.group(1).strip()
                title = match.group(2).strip()
                start_page = int(match.group(3))
                end_page = int(match.group(4))

                paragraphs.append({
                    'id': para_id,
                    'title': title,
                    'start_page': start_page,
                    'end_page': end_page
                })

    return paragraphs

def create_docx_from_images(paragraph_info):
    """
    Create a .docx document from a group of images.

    Args:
        paragraph_info: A dictionary containing paragraph information
            - id: The paragraph ID (number or code)
            - title: The title of the paragraph
            - start_page: The starting page number
            - end_page: The ending page number

    Returns:
        The path to the created document
    """
    para_id = paragraph_info['id']
    title = paragraph_info['title']
    start_page = paragraph_info['start_page']
    end_page = paragraph_info['end_page']

    # Create a new document
    doc = Document()
    doc.add_heading(f"{para_id}. {title}", level=1)

    # Create a temporary directory for image processing
    with tempfile.TemporaryDirectory() as temp_dir:
        # Add each page as an image
        for page_num in range(start_page, end_page + 1):
            # Calculate the image file number (with offset)
            image_num = page_num + PAGE_OFFSET
            image_path = os.path.join(IMAGES_DIR, f"page_{image_num:03d}.png")

            if os.path.exists(image_path):
                # Open the image to get its dimensions
                with Image.open(image_path) as img:
                    width, height = img.size

                # Add page marker before the image
                doc.add_paragraph(f"[[page::{page_num}]]")

                # Add the image to the document
                # Scale the image to fit within the page width
                doc.add_picture(image_path, width=Inches(6))
            else:
                print(f"Warning: Image for page {page_num} (file: {image_path}) not found.")

        # Create the filename based on the paragraph ID and title
        if para_id.isdigit() or '-' in para_id:
            # For numbered paragraphs (e.g., "1", "32-33")
            filename = f"para_{para_id}_{title}"
        else:
            # For special sections (e.g., "intro", "chapter_I")
            filename = f"{para_id}_{title}"

        # Clean up the filename
        safe_filename = re.sub(r'[^\w\s-]', '', filename)
        safe_filename = re.sub(r'[\s-]+', '_', safe_filename)
        safe_filename = safe_filename.lstrip('_')

        output_path = os.path.join(OUTPUT_DIR, f"{safe_filename}.docx")

        # Save the document
        doc.save(output_path)

    return output_path

def main():
    """Main function to orchestrate the process."""
    # Clean up the output directory
    if os.path.exists(OUTPUT_DIR):
        print(f"Cleaning up output directory: {OUTPUT_DIR}")
        for file in os.listdir(OUTPUT_DIR):
            file_path = os.path.join(OUTPUT_DIR, file)
            if os.path.isfile(file_path):
                os.remove(file_path)

    print("Reading paragraphs list...")
    paragraphs = read_paragraphs_list()

    print(f"Found {len(paragraphs)} paragraphs.")

    # Process each paragraph
    for i, paragraph in enumerate(paragraphs):
        print(f"Processing {i+1}/{len(paragraphs)}: {paragraph['id']} - {paragraph['title']}")
        print(f"  Pages: {paragraph['start_page']} to {paragraph['end_page']}")

        # Create a .docx document for this paragraph
        output_path = create_docx_from_images(paragraph)
        print(f"  Created: {output_path}")

if __name__ == "__main__":
    main()
